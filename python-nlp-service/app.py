"""
🤖 NLP Microservice for CV Analysis
Phase 2 - Job Application Manager

Features:
- PDF text extraction (pdfplumber)
- Word document extraction (python-docx)
- NLP analysis with Spacy
- Skills extraction
- Experience detection
- Education parsing

Author: AI Team
Version: 2.0
"""

from flask import Flask, request, jsonify
from flask_cors import CORS
import spacy
import pdfplumber
from docx import Document
import base64
import io
import re
from typing import Dict, List, Any

app = Flask(__name__)
CORS(app)  # Enable CORS for Spring Boot communication

# Load Spacy English model
print("📦 Loading Spacy model...")
try:
    nlp = spacy.load("en_core_web_sm")
    print("✅ Spacy model loaded successfully!")
except OSError:
    print("❌ Spacy model not found. Please run: python -m spacy download en_core_web_sm")
    nlp = None

# Common technical skills database
TECH_SKILLS = {
    # Programming Languages
    'python', 'java', 'javascript', 'typescript', 'c++', 'c#', 'ruby', 'php', 'swift', 'kotlin',
    'go', 'rust', 'scala', 'r', 'matlab', 'perl', 'bash', 'shell',
    
    # Frontend
    'react', 'reactjs', 'angular', 'vue', 'vuejs', 'html', 'css', 'sass', 'less',
    'jquery', 'bootstrap', 'tailwind', 'webpack', 'redux', 'nextjs', 'nuxtjs',
    
    # Backend
    'spring', 'spring boot', 'django', 'flask', 'express', 'nodejs', 'node.js',
    'fastapi', 'laravel', 'rails', 'asp.net', '.net', 'hibernate',
    
    # Databases
    'mysql', 'postgresql', 'mongodb', 'redis', 'cassandra', 'oracle', 'sql server',
    'dynamodb', 'elasticsearch', 'neo4j', 'sqlite',
    
    # Cloud & DevOps
    'aws', 'azure', 'gcp', 'google cloud', 'docker', 'kubernetes', 'jenkins',
    'gitlab', 'github', 'terraform', 'ansible', 'ci/cd', 'circleci',
    
    # Data Science & AI
    'machine learning', 'deep learning', 'tensorflow', 'pytorch', 'scikit-learn',
    'pandas', 'numpy', 'keras', 'nlp', 'computer vision', 'data science',
    
    # Other
    'git', 'agile', 'scrum', 'rest api', 'graphql', 'microservices', 'jira',
    'linux', 'unix', 'windows', 'macos'
}

@app.route('/')
def home():
    """Health check endpoint"""
    return jsonify({
        'status': 'online',
        'service': 'NLP CV Analysis',
        'version': '2.0',
        'spacy_loaded': nlp is not None
    })

@app.route('/api/extract-cv-data', methods=['POST'])
def extract_cv_data():
    """
    Main endpoint: Extract and analyze CV data
    
    Request JSON:
    {
        "resume": "base64_encoded_file",
        "fileType": "pdf" or "docx"
    }
    
    Response JSON:
    {
        "success": true,
        "skills": [...],
        "experience_years": 5,
        "education": [...],
        "raw_text": "...",
        "word_count": 500
    }
    """
    try:
        print("\n" + "="*50)
        print("🤖 NLP SERVICE - New Request Received")
        print("="*50)
        
        # Get data from request
        data = request.json
        
        if not data:
            return jsonify({'success': False, 'error': 'No JSON data provided'}), 400
        
        resume_base64 = data.get('resume')
        file_type = data.get('fileType', 'pdf').lower()
        
        if not resume_base64:
            return jsonify({'success': False, 'error': 'No resume provided'}), 400
        
        print(f"📄 File Type: {file_type}")
        print(f"📦 Base64 Length: {len(resume_base64)} chars")
        
        # Decode Base64
        try:
            # Handle data URL format (data:application/pdf;base64,...)
            if ',' in resume_base64:
                resume_base64 = resume_base64.split(',')[1]
            
            cv_bytes = base64.b64decode(resume_base64)
            print(f"✅ Decoded to {len(cv_bytes)} bytes")
        except Exception as e:
            return jsonify({'success': False, 'error': f'Base64 decode error: {str(e)}'}), 400
        
        # Extract text based on file type
        text = ""
        if file_type in ['pdf']:
            text = extract_text_from_pdf(cv_bytes)
        elif file_type in ['doc', 'docx']:
            text = extract_text_from_doc(cv_bytes)
        else:
            return jsonify({'success': False, 'error': f'Unsupported file type: {file_type}'}), 400
        
        print(f"📝 Extracted {len(text)} characters of text")
        print(f"📊 Word count: {len(text.split())}")
        
        if not text or len(text) < 50:
            return jsonify({
                'success': False,
                'error': 'Could not extract sufficient text from document'
            }), 400
        
        # NLP Analysis
        if nlp:
            print("🧠 Starting NLP analysis...")
            doc = nlp(text[:1000000])  # Limit to 1M chars for Spacy
            
            # Extract information
            skills = extract_skills(doc, text)
            experience_years = extract_experience_years(text)
            education = extract_education(doc, text)
            
            print(f"🔧 Skills found: {len(skills)}")
            print(f"📅 Experience: {experience_years} years")
            print(f"🎓 Education: {len(education)} entries")
        else:
            # Fallback without Spacy
            skills = extract_skills_simple(text)
            experience_years = extract_experience_years(text)
            education = []
        
        # Prepare response
        response = {
            'success': True,
            'skills': skills,
            'experience_years': experience_years,
            'education': education,
            'raw_text': text[:1000],  # First 1000 chars for debugging
            'full_text_length': len(text),
            'word_count': len(text.split())
        }
        
        print("✅ Analysis complete!")
        print("="*50 + "\n")
        
        return jsonify(response)
        
    except Exception as e:
        print(f"❌ Error: {str(e)}")
        import traceback
        traceback.print_exc()
        return jsonify({
            'success': False,
            'error': str(e),
            'type': type(e).__name__
        }), 500

def extract_text_from_pdf(pdf_bytes: bytes) -> str:
    """Extract text from PDF using pdfplumber"""
    try:
        text = ""
        with pdfplumber.open(io.BytesIO(pdf_bytes)) as pdf:
            print(f"📄 PDF has {len(pdf.pages)} pages")
            for i, page in enumerate(pdf.pages):
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
                print(f"   Page {i+1}: {len(page_text) if page_text else 0} chars")
        return text.strip()
    except Exception as e:
        print(f"❌ PDF extraction error: {e}")
        raise Exception(f"Failed to extract PDF: {str(e)}")

def extract_text_from_doc(doc_bytes: bytes) -> str:
    """Extract text from Word document"""
    try:
        doc = Document(io.BytesIO(doc_bytes))
        text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
        print(f"📄 Document has {len(doc.paragraphs)} paragraphs")
        return text.strip()
    except Exception as e:
        print(f"❌ DOC extraction error: {e}")
        raise Exception(f"Failed to extract DOC: {str(e)}")

def extract_skills(doc, text: str) -> List[str]:
    """Extract skills using Spacy NER and keyword matching"""
    skills_found = set()
    text_lower = text.lower()
    
    # Method 1: Match against known tech skills
    for skill in TECH_SKILLS:
        # Use word boundaries to avoid false matches
        pattern = r'\b' + re.escape(skill) + r'\b'
        if re.search(pattern, text_lower):
            skills_found.add(skill)
    
    # Method 2: Extract from common sections
    skills_section = extract_skills_section(text)
    if skills_section:
        section_skills = extract_skills_from_section(skills_section)
        skills_found.update(section_skills)
    
    return sorted(list(skills_found))

def extract_skills_simple(text: str) -> List[str]:
    """Simple skill extraction without Spacy (fallback)"""
    skills_found = set()
    text_lower = text.lower()
    
    for skill in TECH_SKILLS:
        if skill in text_lower:
            skills_found.add(skill)
    
    return sorted(list(skills_found))

def extract_skills_section(text: str) -> str:
    """Extract the Skills section from CV"""
    # Common section headers
    patterns = [
        r'(?i)skills?\s*:?\s*\n(.*?)(?=\n\n|\n[A-Z]|$)',
        r'(?i)technical skills?\s*:?\s*\n(.*?)(?=\n\n|\n[A-Z]|$)',
        r'(?i)competencies\s*:?\s*\n(.*?)(?=\n\n|\n[A-Z]|$)'
    ]
    
    for pattern in patterns:
        match = re.search(pattern, text, re.DOTALL)
        if match:
            return match.group(1)
    
    return ""

def extract_skills_from_section(section: str) -> List[str]:
    """Extract skills from a skills section"""
    skills_found = set()
    
    # Split by common delimiters
    items = re.split(r'[,;•\n]', section)
    
    for item in items:
        item = item.strip().lower()
        if item in TECH_SKILLS:
            skills_found.add(item)
        
        # Check for multi-word skills
        for skill in TECH_SKILLS:
            if skill in item:
                skills_found.add(skill)
    
    return list(skills_found)

def extract_experience_years(text: str) -> int:
    """Extract years of experience using regex patterns"""
    patterns = [
        r'(\d+)\+?\s*years?\s*(?:of)?\s*experience',
        r'experience:\s*(\d+)\+?\s*years?',
        r'(\d+)-\d+\s*years?\s*(?:of)?\s*experience',
        r'over\s*(\d+)\s*years',
        r'more than\s*(\d+)\s*years'
    ]
    
    for pattern in patterns:
        match = re.search(pattern, text, re.IGNORECASE)
        if match:
            years = int(match.group(1))
            print(f"   📅 Found experience: {years} years")
            return years
    
    # Alternative: Count work experience entries
    experience_count = count_work_experiences(text)
    if experience_count > 0:
        # Estimate: 2 years per job position
        estimated_years = experience_count * 2
        print(f"   📅 Estimated from {experience_count} positions: {estimated_years} years")
        return estimated_years
    
    return 0

def count_work_experiences(text: str) -> int:
    """Count number of work experience entries"""
    # Look for date patterns like "2020-2023", "Jan 2020 - Present"
    date_patterns = [
        r'\d{4}\s*[-–]\s*\d{4}',
        r'\d{4}\s*[-–]\s*(?:present|current)',
        r'(?:jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec)[a-z]*\s*\d{4}'
    ]
    
    count = 0
    for pattern in date_patterns:
        matches = re.findall(pattern, text, re.IGNORECASE)
        count += len(matches)
    
    return min(count, 10)  # Cap at 10 positions

def extract_education(doc, text: str) -> List[str]:
    """Extract education information"""
    education = []
    
    # Degrees to look for
    DEGREES = [
        'bachelor', 'bsc', 'b.sc', 'ba', 'b.a',
        'master', 'msc', 'm.sc', 'ma', 'm.a', 'mba',
        'phd', 'ph.d', 'doctorate',
        'diploma', 'associate', 'certification'
    ]
    
    text_lower = text.lower()
    
    # Extract education section
    edu_section = extract_education_section(text)
    if edu_section:
        # Look for degrees in education section
        for line in edu_section.split('\n'):
            line = line.strip()
            if len(line) > 10:  # Minimum length
                for degree in DEGREES:
                    if degree in line.lower():
                        education.append(line[:200])  # Limit length
                        break
    
    # Fallback: Search entire text
    if not education:
        for degree in DEGREES:
            pattern = rf'(?i)\b{degree}\b[^\n]{{0,100}}'
            matches = re.findall(pattern, text)
            education.extend(matches[:3])  # Max 3 entries
    
    return list(set(education))[:5]  # Remove duplicates, max 5

def extract_education_section(text: str) -> str:
    """Extract the Education section from CV"""
    patterns = [
        r'(?i)education\s*:?\s*\n(.*?)(?=\n\n[A-Z]|$)',
        r'(?i)academic\s*:?\s*\n(.*?)(?=\n\n[A-Z]|$)',
        r'(?i)qualifications\s*:?\s*\n(.*?)(?=\n\n[A-Z]|$)'
    ]
    
    for pattern in patterns:
        match = re.search(pattern, text, re.DOTALL)
        if match:
            return match.group(1)[:1000]  # Limit to 1000 chars
    
    return ""

if __name__ == '__main__':
    print("\n" + "="*50)
    print("🚀 Starting NLP Microservice")
    print("="*50)
    print("📍 Endpoint: http://localhost:5000")
    print("📌 API: POST /api/extract-cv-data")
    print("="*50 + "\n")
    
    app.run(host='0.0.0.0', port=5000, debug=True)
